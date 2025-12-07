"""
Document Processing Service
Handles document parsing, chunking, and embedding generation
"""
import logging
import hashlib
import os
from typing import List, Dict, Any, Optional, BinaryIO
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import tiktoken

from app.adapters.vectorstore.base import Document
from app.adapters.vectorstore import get_vectorstore_adapter
from app.adapters.embedding import get_embedding_adapter
from app.schemas.document import (
    DocumentMetadata,
    DocumentChunk,
    DocumentType,
    DocumentUploadResponse
)
from app.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Service for processing documents into searchable chunks.
    
    Pipeline:
    1. Parse document (PDF/Word/Markdown)
    2. Split into chunks
    3. Generate embeddings
    4. Store in vector database
    """
    
    def __init__(self):
        self.vectorstore = get_vectorstore_adapter()
        self.embedding = get_embedding_adapter()
        self.chunk_size = getattr(settings, 'CHUNK_SIZE', 1000)
        self.chunk_overlap = getattr(settings, 'CHUNK_OVERLAP', 200)
        
        # Initialize tokenizer for accurate chunking
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        except Exception:
            self.tokenizer = None
            logger.warning("Tiktoken not available, using character-based chunking")
    
    async def initialize(self):
        """Initialize the vector store connection."""
        await self.vectorstore.initialize()
        logger.info("Document processor initialized")
    
    async def process_document(
        self,
        file: BinaryIO,
        filename: str,
        metadata: Dict[str, Any] = None,
        collection_name: str = "default"
    ) -> DocumentUploadResponse:
        """
        Process a document file and store in vector database.
        
        Args:
            file: File-like object containing document data
            filename: Original filename
            metadata: Additional metadata (tags, description, source)
            collection_name: Target collection name
            
        Returns:
            DocumentUploadResponse with processing results
        """
        start_time = datetime.now()
        metadata = metadata or {}
        
        try:
            # Determine document type
            doc_type = self._get_document_type(filename)
            if not doc_type:
                return DocumentUploadResponse(
                    status="failed",
                    message=f"Unsupported file type: {filename}",
                    chunks_created=0,
                    processing_time_seconds=0
                )
            
            # Read file content
            file_content = file.read()
            file_size = len(file_content)
            
            # Check file size
            max_size = getattr(settings, 'MAX_FILE_SIZE_MB', 50) * 1024 * 1024
            if file_size > max_size:
                return DocumentUploadResponse(
                    status="failed",
                    message=f"File too large: {file_size / 1024 / 1024:.2f} MB (max: {max_size / 1024 / 1024} MB)",
                    chunks_created=0,
                    processing_time_seconds=0
                )
            
            # Generate document ID
            doc_id = self._generate_document_id(filename, file_content)
            
            # Parse document
            logger.info(f"Parsing document: {filename} ({doc_type})")
            parsed_content, page_count = await self._parse_document(
                file_content, doc_type
            )
            
            if not parsed_content:
                return DocumentUploadResponse(
                    status="failed",
                    message="Failed to extract text from document",
                    chunks_created=0,
                    processing_time_seconds=0
                )
            
            # Split into chunks
            logger.info(f"Splitting document into chunks (size={self.chunk_size})")
            chunks = self._split_into_chunks(parsed_content, doc_id)
            
            if not chunks:
                return DocumentUploadResponse(
                    status="failed",
                    message="No content chunks created",
                    chunks_created=0,
                    processing_time_seconds=0
                )
            
            # Generate embeddings
            logger.info(f"Generating embeddings for {len(chunks)} chunks")
            chunk_texts = [chunk.content for chunk in chunks]
            embeddings = await self.embedding.embed_texts(chunk_texts)
            
            # Create Document objects for vector store
            documents = []
            
            # Convert tags list to comma-separated string (Chroma doesn't accept lists)
            tags_str = ",".join(metadata.get("tags", []))
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                # Handle None values - Chroma requires str, int, float, bool, or None
                chunk_metadata = {
                    "document_id": doc_id,
                    "filename": filename,
                    "chunk_index": chunk.chunk_index,
                    "tags": tags_str,  # String instead of list
                    "source": metadata.get("source", "") or "",
                    "description": metadata.get("description", "") or "",
                    "uploaded_at": datetime.now().isoformat()
                }
                
                # Only add page_number and section_title if they exist
                if chunk.page_number is not None:
                    chunk_metadata["page_number"] = chunk.page_number
                if chunk.section_title is not None:
                    chunk_metadata["section_title"] = chunk.section_title
                
                doc = Document(
                    id=chunk.chunk_id,
                    content=chunk.content,
                    metadata=chunk_metadata,
                    embedding=embedding
                )
                documents.append(doc)
            
            # Store in vector database
            logger.info(f"Storing {len(documents)} chunks in vector store")
            await self.vectorstore.add_documents(documents, collection_name)
            
            # Create document metadata
            doc_metadata = DocumentMetadata(
                document_id=doc_id,
                filename=filename,
                document_type=doc_type,
                file_size=file_size,
                total_chunks=len(chunks),
                total_pages=page_count,
                total_characters=len(parsed_content),
                uploaded_at=start_time,
                processed_at=datetime.now(),
                tags=metadata.get("tags", []),
                description=metadata.get("description"),
                source=metadata.get("source")
            )
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            logger.info(f"Document processed successfully: {doc_id} ({len(chunks)} chunks)")
            
            return DocumentUploadResponse(
                status="success",
                message=f"Document processed successfully",
                document=doc_metadata,
                chunks_created=len(chunks),
                processing_time_seconds=processing_time
            )
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}", exc_info=True)
            processing_time = (datetime.now() - start_time).total_seconds()
            return DocumentUploadResponse(
                status="failed",
                message=f"Processing error: {str(e)}",
                chunks_created=0,
                processing_time_seconds=processing_time
            )
    
    def _get_document_type(self, filename: str) -> Optional[DocumentType]:
        """Determine document type from filename."""
        ext = filename.lower().split('.')[-1]
        type_map = {
            'pdf': DocumentType.PDF,
            'docx': DocumentType.WORD,
            'doc': DocumentType.WORD,
            'md': DocumentType.MARKDOWN,
            'markdown': DocumentType.MARKDOWN,
            'txt': DocumentType.TEXT
        }
        return type_map.get(ext)
    
    def _generate_document_id(self, filename: str, content: bytes) -> str:
        """Generate unique document ID based on filename and content hash."""
        content_hash = hashlib.md5(content).hexdigest()[:8]
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        safe_name = "".join(c for c in filename if c.isalnum() or c in "._-")[:20]
        return f"doc_{timestamp}_{safe_name}_{content_hash}"
    
    async def _parse_document(
        self, 
        content: bytes, 
        doc_type: DocumentType
    ) -> tuple[str, Optional[int]]:
        """
        Parse document and extract text.
        
        Returns:
            Tuple of (extracted_text, page_count)
        """
        if doc_type == DocumentType.PDF:
            return self._parse_pdf(content)
        elif doc_type == DocumentType.WORD:
            return self._parse_word(content)
        elif doc_type in (DocumentType.MARKDOWN, DocumentType.TEXT):
            return self._parse_text(content)
        else:
            return "", None
    
    def _parse_pdf(self, content: bytes) -> tuple[str, int]:
        """Extract text from PDF using PyMuPDF."""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    # Add page marker for reference
                    text_parts.append(f"[PAGE {page_num + 1}]\n{text}")
            
            doc.close()
            return "\n\n".join(text_parts), len(doc)
            
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return "", 0
    
    def _parse_word(self, content: bytes) -> tuple[str, None]:
        """Extract text from Word document."""
        try:
            import io
            doc = DocxDocument(io.BytesIO(content))
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check if it's a heading
                    if para.style.name.startswith('Heading'):
                        text_parts.append(f"\n## {para.text}\n")
                    else:
                        text_parts.append(para.text)
            
            return "\n".join(text_parts), None
            
        except Exception as e:
            logger.error(f"Word parsing error: {e}")
            return "", None
    
    def _parse_text(self, content: bytes) -> tuple[str, None]:
        """Parse plain text or markdown."""
        try:
            # Try UTF-8 first, then fall back to other encodings
            for encoding in ['utf-8', 'utf-16', 'latin-1']:
                try:
                    return content.decode(encoding), None
                except UnicodeDecodeError:
                    continue
            return content.decode('utf-8', errors='ignore'), None
        except Exception as e:
            logger.error(f"Text parsing error: {e}")
            return "", None
    
    def _split_into_chunks(
        self, 
        text: str, 
        document_id: str
    ) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks.
        
        Uses sentence-aware splitting to avoid breaking mid-sentence.
        """
        chunks = []
        
        # Clean text
        text = text.strip()
        if not text:
            return chunks
        
        # Split by paragraphs first, then recombine to target size
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        current_page = 1
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Check for page markers
            if para.startswith('[PAGE '):
                try:
                    page_end = para.index(']')
                    current_page = int(para[6:page_end])
                    para = para[page_end + 1:].strip()
                except (ValueError, IndexError):
                    pass
            
            # Check if adding this paragraph exceeds chunk size
            test_chunk = current_chunk + "\n\n" + para if current_chunk else para
            
            if len(test_chunk) > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append(DocumentChunk(
                    chunk_id=f"{document_id}_chunk_{chunk_index}",
                    document_id=document_id,
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    page_number=current_page
                ))
                chunk_index += 1
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + "\n\n" + para if overlap_text else para
            else:
                current_chunk = test_chunk
        
        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                chunk_id=f"{document_id}_chunk_{chunk_index}",
                document_id=document_id,
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                page_number=current_page
            ))
        
        return chunks
    
    def _get_overlap_text(self, text: str) -> str:
        """Get the last N characters for overlap."""
        if len(text) <= self.chunk_overlap:
            return text
        
        # Try to break at sentence boundary
        overlap_region = text[-self.chunk_overlap:]
        
        # Find last sentence end in overlap region
        for sep in ['. ', '。', '! ', '? ']:
            last_sep = overlap_region.rfind(sep)
            if last_sep != -1:
                return overlap_region[last_sep + len(sep):]
        
        return overlap_region
    
    async def delete_document(
        self, 
        document_id: str, 
        collection_name: str = "default"
    ) -> bool:
        """
        Delete a document and all its chunks.
        
        Args:
            document_id: The document ID to delete
            collection_name: Collection name
            
        Returns:
            True if deleted successfully
        """
        try:
            # Get all chunk IDs for this document
            # This is a simplified version - in production you'd query by metadata
            stats = await self.vectorstore.get_collection_stats(collection_name)
            
            # For now, we'll delete by pattern matching chunk IDs
            # In a real implementation, you'd use metadata filtering
            logger.info(f"Deleting document: {document_id}")
            
            # This is a placeholder - actual implementation depends on vector store
            # capability to query by metadata
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document: {e}")
            return False
    
    async def list_documents(
        self, 
        collection_name: str = "default"
    ) -> List[DocumentMetadata]:
        """
        List all documents in a collection.
        
        Note: This is a simplified implementation. In production,
        you'd store document metadata in a separate database.
        """
        # This would typically query a metadata store
        # For now, return collection stats
        stats = await self.vectorstore.get_collection_stats(collection_name)
        return []
    
    def get_processor_info(self) -> Dict[str, Any]:
        """Get information about the document processor configuration."""
        return {
            "vectorstore": self.vectorstore.get_store_name(),
            "embedding_provider": self.embedding.get_provider_name(),
            "embedding_model": self.embedding.get_model_name(),
            "embedding_dimension": self.embedding.get_embedding_dimension(),
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap
        }